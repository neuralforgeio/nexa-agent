"""Read a .docx file: paragraphs then tables. Uses python-docx."""
from __future__ import annotations

from typing import Any

from tools._paths import resolve_in_workspace

try:
    import docx  # python-docx
except Exception:  # pragma: no cover
    docx = None  # type: ignore[assignment]

READ_DOCX_SCHEMA = {
    "type": "object",
    "properties": {
        "path": {"type": "string", "description": "Workspace-relative path to the .docx file."},
    },
    "required": ["path"],
}


async def read_docx(path: str, **_: Any) -> str:
    """Extract all paragraphs, then each table as pipe-separated rows."""
    if docx is None:
        return "read_docx requires the 'python-docx' package (pip install python-docx)."
    full = resolve_in_workspace(path)
    if not full.exists():
        raise ValueError(f"file not found: '{path}'")
    document = docx.Document(str(full))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for ti, table in enumerate(document.tables, 1):
        parts.append(f"--- table {ti} ---")
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts) or "(empty document)"
